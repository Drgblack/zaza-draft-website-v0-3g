from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# ENGLISH PARENT EMAIL TEMPLATES
# ============================================================================
print("Creating parent-email-templates.docx (English)...")

doc = Document()

# Title
title = doc.add_heading('Parent Email Templates', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(139, 92, 246)

subtitle = doc.add_paragraph('Ready-to-Use Templates for AI Communication with Parents')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

doc.add_paragraph()

# Introduction
doc.add_heading('Introduction', 1)
doc.add_paragraph(
    'Effective communication with parents about AI in the classroom is essential for building '
    'trust and ensuring student success. These templates are designed to help you communicate '
    'clearly, professionally, and proactively about your use of AI tools.'
)

# Template 1
doc.add_heading('Template 1: Introduction to AI in the Classroom', 1)

p = doc.add_paragraph()
p.add_run('Subject: ').bold = True
p.add_run('Exciting Updates About AI-Enhanced Learning in [Grade/Class]')

doc.add_paragraph()
doc.add_paragraph('Dear [Parent/Guardian Name],')
doc.add_paragraph()
doc.add_paragraph(
    "I hope this message finds you well. I'm writing to share some exciting updates about "
    "how we're enhancing learning in [your classroom/grade level] this year."
)
doc.add_paragraph()
doc.add_paragraph("We'll be thoughtfully integrating AI tools to support your child's learning. These tools will help us:")

benefits = [
    'Provide more personalized learning experiences',
    'Give faster, more detailed feedback on assignments',
    'Create engaging, differentiated activities',
    'Save time on administrative tasks so I can focus more on teaching'
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('What This Means for Your Child:').bold = True
doc.add_paragraph(
    'Your child will still receive the same personal attention and direct instruction. '
    'AI is simply a tool that helps me be more effective and responsive to each student\'s needs. '
    'All AI-generated content is reviewed by me before use, and student data privacy remains our top priority.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Your Role:').bold = True
doc.add_paragraph(
    "I encourage you to ask your child about their learning experiences and any new tools they're using. "
    "If you have any questions or concerns about AI in the classroom, please don't hesitate to reach out."
)

doc.add_paragraph()
doc.add_paragraph(
    "I'm excited about the possibilities this brings for enhancing your child's education, "
    "and I look forward to keeping you updated on their progress."
)

doc.add_paragraph()
doc.add_paragraph('Warm regards,')
doc.add_paragraph('[Your Name]')
doc.add_paragraph('[Your Title]')
doc.add_paragraph('[Contact Information]')

doc.add_page_break()

# Template 2
doc.add_heading('Template 2: Monthly Progress Update', 1)

p = doc.add_paragraph()
p.add_run('Subject: ').bold = True
p.add_run('[Student Name]\'s Progress Update - [Month]')

doc.add_paragraph()
doc.add_paragraph('Dear [Parent/Guardian Name],')
doc.add_paragraph()
doc.add_paragraph("I wanted to share [Student Name]'s progress this month and highlight some wonderful achievements.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Academic Highlights:').bold = True
highlights = [
    '[Specific achievement or improvement area]',
    '[Growth observed in particular subject/skill]',
    '[Positive participation or effort noted]'
]
for highlight in highlights:
    doc.add_paragraph(highlight, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('How AI Tools Are Supporting Learning:').bold = True
doc.add_paragraph(
    'This month, [Student Name] has been using AI-assisted tools for [specific activity]. '
    'These tools have helped by [specific benefit].'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Areas of Focus Moving Forward:').bold = True
doc.add_paragraph('[Describe 1-2 areas where continued growth is expected or needed]')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('How You Can Support at Home:').bold = True
home_support = [
    '[Specific suggestion related to current curriculum]',
    '[Activity or practice that would reinforce classroom learning]'
]
for item in home_support:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph("Please let me know if you'd like to schedule a call to discuss [Student Name]'s progress in more detail.")

doc.add_paragraph()
doc.add_paragraph('Best regards,')
doc.add_paragraph('[Your Name]')
doc.add_paragraph('[Contact Information]')

doc.add_page_break()

# Template 3
doc.add_heading('Template 3: Addressing Parent Concerns', 1)

p = doc.add_paragraph()
p.add_run('Subject: ').bold = True
p.add_run('Re: Your Questions About AI in Our Classroom')

doc.add_paragraph()
doc.add_paragraph('Dear [Parent/Guardian Name],')
doc.add_paragraph()
doc.add_paragraph(
    'Thank you for reaching out with your questions and concerns about AI use in our classroom. '
    'I appreciate your engagement and want to address your thoughts directly.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Your Concern:').bold = True
doc.add_paragraph('[Summarize the parent\'s specific concern]')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('My Response:').bold = True
doc.add_paragraph(
    '[Provide detailed, empathetic response addressing the concern. '
    'Include specific examples of how you\'re mitigating the issue they raised]'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('What I\'m Doing to Ensure [Student Safety/Quality/Privacy]:').bold = True
measures = [
    '[Specific measure 1]',
    '[Specific measure 2]',
    '[Specific measure 3]'
]
for measure in measures:
    doc.add_paragraph(measure, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Moving Forward:').bold = True
doc.add_paragraph('[Explain any adjustments you\'re willing to make or alternatives you can offer]')

doc.add_paragraph()
doc.add_paragraph(
    'I value your perspective and want to ensure [Student Name] has the best possible learning experience. '
    'Would you be available for a phone call or in-person meeting to discuss this further?'
)

doc.add_paragraph()
doc.add_paragraph('Sincerely,')
doc.add_paragraph('[Your Name]')
doc.add_paragraph('[Contact Information]')

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('Created by Zaza Draft - AI Tools for Teachers')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.save('public/downloads/parent-email-templates.docx')
print("   Done!\n")

# ============================================================================
# GERMAN PARENT EMAIL TEMPLATES
# ============================================================================
print("Creating eltern-email-vorlagen.docx (German)...")

doc = Document()

# Title
title = doc.add_heading('Eltern-E-Mail-Vorlagen', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(139, 92, 246)

subtitle = doc.add_paragraph('Gebrauchsfertige Vorlagen für die KI-Kommunikation mit Eltern')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

doc.add_paragraph()

# Introduction
doc.add_heading('Einführung', 1)
doc.add_paragraph(
    'Eine effektive Kommunikation mit Eltern über KI im Klassenzimmer ist entscheidend für den Aufbau '
    'von Vertrauen und den Erfolg der Schüler. Diese Vorlagen helfen Ihnen, klar, professionell und '
    'proaktiv über Ihren Einsatz von KI-Tools zu kommunizieren.'
)

# Template 1
doc.add_heading('Vorlage 1: Einführung der KI im Klassenzimmer', 1)

p = doc.add_paragraph()
p.add_run('Betreff: ').bold = True
p.add_run('Spannende Neuigkeiten zum KI-unterstützten Lernen in [Klasse/Jahrgang]')

doc.add_paragraph()
doc.add_paragraph('Sehr geehrte/r [Name des Erziehungsberechtigten],')
doc.add_paragraph()
doc.add_paragraph(
    'ich möchte Ihnen einige spannende Neuigkeiten mitteilen, wie wir das Lernen in '
    '[Ihrer Klasse/Jahrgangsstufe] in diesem Jahr verbessern.'
)
doc.add_paragraph()
doc.add_paragraph('Wir werden KI-Tools durchdacht integrieren, um das Lernen Ihres Kindes zu unterstützen. Diese Tools helfen uns:')

benefits_de = [
    'Personalisiertere Lernerfahrungen bereitzustellen',
    'Schnelleres, detaillierteres Feedback zu Aufgaben zu geben',
    'Ansprechende, differenzierte Aktivitäten zu erstellen',
    'Zeit bei administrativen Aufgaben zu sparen, damit ich mich mehr auf das Unterrichten konzentrieren kann'
]
for benefit in benefits_de:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Was das für Ihr Kind bedeutet:').bold = True
doc.add_paragraph(
    'Ihr Kind wird weiterhin die gleiche persönliche Aufmerksamkeit und direkte Anleitung erhalten. '
    'KI ist einfach ein Werkzeug, das mir hilft, effektiver zu sein und auf die Bedürfnisse jedes '
    'Schülers einzugehen. Alle KI-generierten Inhalte werden von mir überprüft, bevor sie verwendet '
    'werden, und der Datenschutz der Schüler hat höchste Priorität.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Ihre Rolle:').bold = True
doc.add_paragraph(
    'Ich ermutige Sie, Ihr Kind nach seinen Lernerfahrungen und neuen Tools zu fragen. '
    'Bei Fragen oder Bedenken zum KI-Einsatz im Klassenzimmer können Sie sich gerne an mich wenden.'
)

doc.add_paragraph()
doc.add_paragraph(
    'Ich freue mich über die Möglichkeiten zur Verbesserung der Bildung Ihres Kindes und '
    'halte Sie gerne über die Fortschritte auf dem Laufenden.'
)

doc.add_paragraph()
doc.add_paragraph('Mit freundlichen Grüßen,')
doc.add_paragraph('[Ihr Name]')
doc.add_paragraph('[Ihre Position]')
doc.add_paragraph('[Kontaktinformationen]')

doc.add_page_break()

# Template 2
doc.add_heading('Vorlage 2: Monatlicher Fortschrittsbericht', 1)

p = doc.add_paragraph()
p.add_run('Betreff: ').bold = True
p.add_run('Fortschrittsbericht für [Schülername] - [Monat]')

doc.add_paragraph()
doc.add_paragraph('Sehr geehrte/r [Name des Erziehungsberechtigten],')
doc.add_paragraph()
doc.add_paragraph('Ich möchte Ihnen den Fortschritt von [Schülername] in diesem Monat mitteilen und einige wunderbare Erfolge hervorheben.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Akademische Höhepunkte:').bold = True
highlights_de = [
    '[Spezifische Leistung oder Verbesserungsbereich]',
    '[Beobachtetes Wachstum in bestimmtem Fach/Fähigkeit]',
    '[Positive Beteiligung oder Anstrengung]'
]
for highlight in highlights_de:
    doc.add_paragraph(highlight, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Wie KI-Tools das Lernen unterstützen:').bold = True
doc.add_paragraph(
    'In diesem Monat hat [Schülername] KI-unterstützte Tools für [spezifische Aktivität] verwendet. '
    'Diese Tools haben geholfen, indem [spezifischer Nutzen].'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Schwerpunkte für die Zukunft:').bold = True
doc.add_paragraph('[Beschreiben Sie 1-2 Bereiche, in denen weiteres Wachstum erwartet wird oder benötigt wird]')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Wie Sie zu Hause unterstützen können:').bold = True
home_support_de = [
    '[Spezifischer Vorschlag zum aktuellen Lehrplan]',
    '[Aktivität oder Übung zur Verstärkung des Klassenzimmer-Lernens]'
]
for item in home_support_de:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Lassen Sie mich wissen, wenn Sie ein Gespräch über die Fortschritte von [Schülername] vereinbaren möchten.')

doc.add_paragraph()
doc.add_paragraph('Mit freundlichen Grüßen,')
doc.add_paragraph('[Ihr Name]')
doc.add_paragraph('[Kontaktinformationen]')

doc.add_page_break()

# Template 3
doc.add_heading('Vorlage 3: Bedenken ansprechen', 1)

p = doc.add_paragraph()
p.add_run('Betreff: ').bold = True
p.add_run('Ihre Fragen zur KI in unserem Klassenzimmer')

doc.add_paragraph()
doc.add_paragraph('Sehr geehrte/r [Name des Erziehungsberechtigten],')
doc.add_paragraph()
doc.add_paragraph(
    'Vielen Dank, dass Sie sich mit Ihren Fragen und Bedenken zur KI-Nutzung in unserem Klassenzimmer '
    'an mich gewandt haben. Ich schätze Ihr Engagement und möchte Ihre Gedanken direkt ansprechen.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Ihre Bedenken:').bold = True
doc.add_paragraph('[Zusammenfassung der spezifischen Bedenken der Eltern]')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Meine Antwort:').bold = True
doc.add_paragraph(
    '[Geben Sie eine detaillierte, einfühlsame Antwort, die die Bedenken anspricht. '
    'Fügen Sie spezifische Beispiele hinzu, wie Sie das angesprochene Problem mindern]'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Was ich tue, um [Sicherheit/Qualität/Datenschutz] zu gewährleisten:').bold = True
measures_de = [
    '[Spezifische Maßnahme 1]',
    '[Spezifische Maßnahme 2]',
    '[Spezifische Maßnahme 3]'
]
for measure in measures_de:
    doc.add_paragraph(measure, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Weiteres Vorgehen:').bold = True
doc.add_paragraph('[Erklären Sie Anpassungen, die Sie vornehmen können, oder Alternativen, die Sie anbieten können]')

doc.add_paragraph()
doc.add_paragraph(
    'Ich schätze Ihre Perspektive und möchte sicherstellen, dass [Schülername] die bestmögliche '
    'Lernerfahrung hat. Wären Sie für ein Telefongespräch oder persönliches Treffen verfügbar, '
    'um dies weiter zu besprechen?'
)

doc.add_paragraph()
doc.add_paragraph('Mit freundlichen Grüßen,')
doc.add_paragraph('[Ihr Name]')
doc.add_paragraph('[Kontaktinformationen]')

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('Erstellt von Zaza Draft - KI-Werkzeuge für Lehrkräfte')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.save('public/downloads/eltern-email-vorlagen.docx')
print("   Done!\n")

print("=" * 70)
print("SUCCESS! Both DOCX files created:")
print("  - parent-email-templates.docx")
print("  - eltern-email-vorlagen.docx")
print("=" * 70)
