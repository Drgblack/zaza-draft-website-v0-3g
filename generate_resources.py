#!/usr/bin/env python3
"""
Zaza Draft Resource Generator
Creates all 14 resources (7 English + 7 German) for the website
Run this script to generate all Excel, Word, and HTML files
"""

import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create output directory
OUTPUT_DIR = "zaza_resources_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

print("=" * 70)
print("ZAZA DRAFT RESOURCE GENERATOR")
print("=" * 70)
print(f"\nGenerating all 14 resources in: {OUTPUT_DIR}/")
print("This will take about 2-3 minutes...\n")

# ============================================================================
# 1. ENGLISH REPORT CARD COMMENTS (Excel)
# ============================================================================
def create_report_card_comments_en():
    print("Creating 1/14: report-card-comments.xlsx (English)...")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Comment Bank"
    
    # Headers
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="8B5CF6", end_color="8B5CF6", fill_type="solid")
    
    headers = ["Subject", "Skill Area", "Tone", "Comment"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Column widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 80
    
    # Comments data
    comments = [
        ("Mathematics", "Achievement", "Positive", "Demonstrates exceptional understanding of mathematical concepts and consistently applies them to solve complex problems."),
        ("Mathematics", "Achievement", "Positive", "Shows strong mastery of grade-level math skills and can explain their reasoning clearly."),
        ("Mathematics", "Problem Solving", "Positive", "Approaches mathematical challenges with confidence and persistence."),
        ("Mathematics", "Effort", "Positive", "Consistently puts forth excellent effort during math lessons and assignments."),
        ("English/Reading", "Comprehension", "Positive", "Demonstrates excellent reading comprehension across various text types and genres."),
        ("English/Reading", "Fluency", "Positive", "Reads fluently with appropriate expression and pacing."),
        ("English/Writing", "Composition", "Positive", "Writes clear, well-organized compositions with strong voice and style."),
        ("English/Writing", "Mechanics", "Positive", "Demonstrates strong command of grammar, spelling, and punctuation."),
        ("Science", "Understanding", "Positive", "Shows excellent curiosity and understanding of scientific concepts."),
        ("Science", "Inquiry Skills", "Positive", "Demonstrates excellent scientific inquiry skills through observations and questions."),
        ("General", "Participation", "Positive", "Actively participates in class discussions with thoughtful contributions."),
        ("General", "Behavior", "Positive", "Demonstrates excellent self-control and respect for classroom expectations."),
        ("General", "Work Habits", "Positive", "Consistently completes assignments on time with careful attention to detail."),
        ("General", "Collaboration", "Positive", "Works exceptionally well with peers and contributes positively to group activities."),
        # Neutral comments
        ("Mathematics", "Achievement", "Neutral", "Shows adequate understanding of mathematical concepts with room for deeper exploration."),
        ("English/Reading", "Comprehension", "Neutral", "Shows basic comprehension but would benefit from deeper text analysis."),
        ("General", "Participation", "Neutral", "Participates occasionally; would benefit from more active engagement."),
        # Needs improvement
        ("Mathematics", "Achievement", "Needs Improvement", "Would benefit from additional practice with fundamental mathematical operations."),
        ("English/Reading", "Comprehension", "Needs Improvement", "Requires support to understand grade-level reading material."),
        ("General", "Work Habits", "Needs Improvement", "Needs to develop stronger study habits and assignment completion skills."),
    ]
    
    for comment in comments:
        ws.append(comment)
    
    # Format cells
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    
    # Info sheet
    info_sheet = wb.create_sheet("How to Use")
    info_sheet.column_dimensions['A'].width = 100
    
    info_content = [
        ["Report Card Comment Bank - User Guide"],
        [""],
        ["This comment bank contains 20+ ready-to-use report card comments organized by:"],
        ["- Subject Area (Mathematics, Reading, Writing, Science, etc.)"],
        ["- Skill Area (Achievement, Comprehension, Problem Solving, etc.)"],
        ["- Tone (Positive, Neutral, Needs Improvement)"],
        [""],
        ["How to Use:"],
        ["1. Filter by Subject or Tone using the dropdown arrows"],
        ["2. Search for keywords using Ctrl+F"],
        ["3. Customize comments to fit your specific students"],
        [""],
        ["Created by Zaza Draft - AI Tools for Teachers"],
        ["Visit zazadraft.com for more resources"],
    ]
    
    for row_content in info_content:
        info_sheet.append(row_content)
    
    info_sheet['A1'].font = Font(bold=True, size=14, color="8B5CF6")
    
    wb.save(f"{OUTPUT_DIR}/report-card-comments.xlsx")
    print("   ✅ Complete\n")

# ============================================================================
# 2. GERMAN REPORT CARD COMMENTS (Excel)
# ============================================================================
def create_report_card_comments_de():
    print("Creating 2/14: zeugniskommentare.xlsx (German)...")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Kommentare"
    
    # Headers
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="8B5CF6", end_color="8B5CF6", fill_type="solid")
    
    headers = ["Fach", "Kompetenzbereich", "Ton", "Kommentar"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 85
    
    # German comments
    comments = [
        ("Mathematik", "Leistung", "Positiv", "Zeigt außergewöhnliches Verständnis mathematischer Konzepte und wendet sie konsequent zur Lösung komplexer Probleme an."),
        ("Mathematik", "Leistung", "Positiv", "Beherrscht die mathematischen Kompetenzen der Jahrgangsstufe sehr gut und kann die Lösungswege klar erklären."),
        ("Mathematik", "Problemlösung", "Positiv", "Geht mathematische Herausforderungen mit Selbstvertrauen und Ausdauer an."),
        ("Mathematik", "Anstrengung", "Positiv", "Zeigt durchgehend ausgezeichnete Anstrengung im Mathematikunterricht und bei Hausaufgaben."),
        ("Deutsch/Lesen", "Leseverstehen", "Positiv", "Zeigt ausgezeichnetes Leseverstehen über verschiedene Textarten und Genres hinweg."),
        ("Deutsch/Lesen", "Leseflüssigkeit", "Positiv", "Liest flüssig mit angemessener Betonung und Geschwindigkeit."),
        ("Deutsch/Schreiben", "Textgestaltung", "Positiv", "Schreibt klare, gut strukturierte Texte mit starkem Ausdruck und Stil."),
        ("Deutsch/Schreiben", "Rechtschreibung", "Positiv", "Zeigt sichere Beherrschung von Grammatik, Rechtschreibung und Zeichensetzung."),
        ("Sachunterricht", "Verständnis", "Positiv", "Zeigt ausgezeichnete Neugier und Verständnis naturwissenschaftlicher Konzepte."),
        ("Sachunterricht", "Forschende Fähigkeiten", "Positiv", "Zeigt ausgezeichnete wissenschaftliche Forschungsfähigkeiten durch Beobachtungen und Fragen."),
        ("Allgemein", "Beteiligung", "Positiv", "Beteiligt sich aktiv an Klassendiskussionen mit durchdachten Beiträgen."),
        ("Allgemein", "Verhalten", "Positiv", "Zeigt ausgezeichnete Selbstkontrolle und Respekt gegenüber Klassenerwartungen."),
        ("Allgemein", "Arbeitsverhalten", "Positiv", "Schließt Aufgaben durchgehend pünktlich mit sorgfältiger Aufmerksamkeit fürs Detail ab."),
        ("Allgemein", "Zusammenarbeit", "Positiv", "Arbeitet außergewöhnlich gut mit Mitschülern und trägt positiv zu Gruppenaktivitäten bei."),
        # Neutral
        ("Mathematik", "Leistung", "Neutral", "Zeigt angemessenes Verständnis mathematischer Konzepte mit Raum für tiefere Auseinandersetzung."),
        ("Deutsch/Lesen", "Leseverstehen", "Neutral", "Zeigt grundlegendes Verständnis, würde jedoch von tieferer Textanalyse profitieren."),
        ("Allgemein", "Beteiligung", "Neutral", "Beteiligt sich gelegentlich; würde von aktiverer Teilnahme profitieren."),
        # Verbesserungsbedarf
        ("Mathematik", "Leistung", "Verbesserungsbedarf", "Würde von zusätzlicher Übung grundlegender mathematischer Operationen profitieren."),
        ("Deutsch/Lesen", "Leseverstehen", "Verbesserungsbedarf", "Benötigt Unterstützung beim Verstehen altersgerechter Lesetexte."),
        ("Allgemein", "Arbeitsverhalten", "Verbesserungsbedarf", "Sollte stärkere Lerngewohnheiten und Fähigkeiten zur Aufgabenerfüllung entwickeln."),
    ]
    
    for comment in comments:
        ws.append(comment)
    
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    
    # Info sheet
    info_sheet = wb.create_sheet("Anleitung")
    info_sheet.column_dimensions['A'].width = 100
    
    info_content = [
        ["Zeugniskommentare - Benutzerhandbuch"],
        [""],
        ["Diese Kommentarsammlung enthält 20+ fertige Zeugniskommentare, organisiert nach:"],
        ["- Fach (Mathematik, Deutsch/Lesen, Sachunterricht, usw.)"],
        ["- Kompetenzbereich (Leistung, Leseverstehen, Problemlösung, usw.)"],
        ["- Ton (Positiv, Neutral, Verbesserungsbedarf)"],
        [""],
        ["Anwendung:"],
        ["1. Filtern Sie nach Fach oder Ton mit den Dropdown-Pfeilen"],
        ["2. Suchen Sie nach Stichwörtern mit Strg+F"],
        ["3. Passen Sie Kommentare an Ihre spezifischen Schüler an"],
        [""],
        ["Erstellt von Zaza Draft - KI-Werkzeuge für Lehrkräfte"],
        ["Besuchen Sie zazadraft.com für weitere Ressourcen"],
    ]
    
    for row_content in info_content:
        info_sheet.append(row_content)
    
    info_sheet['A1'].font = Font(bold=True, size=14, color="8B5CF6")
    
    wb.save(f"{OUTPUT_DIR}/zeugniskommentare.xlsx")
    print("   ✅ Complete\n")

# ============================================================================
# 3. ENGLISH IEP TEMPLATES (Word)
# ============================================================================
def create_iep_templates_en():
    print("Creating 3/14: iep-goal-templates.docx (English)...")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('IEP Goal Templates', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(139, 92, 246)
    
    subtitle = doc.add_paragraph('SMART Goals for Special Education')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    # Introduction
    doc.add_heading('Understanding SMART IEP Goals', 1)
    doc.add_paragraph(
        'Effective IEP goals are SMART: Specific, Measurable, Achievable, Relevant, and Time-bound.'
    )
    
    # Template
    doc.add_heading('Universal Goal Template', 2)
    template = doc.add_paragraph()
    template.add_run('By ').italic = True
    template.add_run('[DATE]').bold = True
    template.add_run(', ').italic = True
    template.add_run('[STUDENT NAME]').bold = True
    template.add_run(' will ').italic = True
    template.add_run('[SPECIFIC ACTION]').bold = True
    template.add_run(' with ').italic = True
    template.add_run('[ACCURACY/SUPPORT]').bold = True
    template.add_run(' as measured by ').italic = True
    template.add_run('[ASSESSMENT]').bold = True
    
    # Sample goals
    doc.add_heading('Sample Goals', 2)
    
    doc.add_heading('Reading Comprehension', 3)
    goals = [
        "By May 2026, [Student] will read grade-level passages and answer comprehension questions with 80% accuracy across 4 out of 5 trials.",
        "By December 2025, [Student] will identify the main idea and three supporting details in informational texts with 85% accuracy.",
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_heading('Mathematics', 3)
    goals = [
        "By June 2026, [Student] will solve multi-step word problems involving addition and subtraction with 80% accuracy across 4 out of 5 opportunities.",
        "By April 2026, [Student] will demonstrate understanding of multiplication facts 0-12 by solving problems with 90% accuracy.",
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Created by Zaza Draft - AI Tools for Teachers')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    
    doc.save(f"{OUTPUT_DIR}/iep-goal-templates.docx")
    print("   ✅ Complete\n")

# ============================================================================
# 4. GERMAN FÖRDERPLAN TEMPLATES (Word)
# ============================================================================
def create_foerderplan_templates_de():
    print("Creating 4/14: foerderplan-vorlagen.docx (German)...")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Förderplan-Vorlagen', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(139, 92, 246)
    
    subtitle = doc.add_paragraph('SMART-Ziele für die Förderplanung')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    # Introduction
    doc.add_heading('SMART-Förderziele verstehen', 1)
    doc.add_paragraph(
        'Effektive Förderziele sind SMART: Spezifisch, Messbar, Attraktiv, Realistisch und Terminiert.'
    )
    
    # Template
    doc.add_heading('Universelle Zielvorlage', 2)
    template = doc.add_paragraph()
    template.add_run('Bis zum ').italic = True
    template.add_run('[DATUM]').bold = True
    template.add_run(' wird ').italic = True
    template.add_run('[SCHÜLERNAME]').bold = True
    template.add_run(' ').italic = True
    template.add_run('[SPEZIFISCHE HANDLUNG]').bold = True
    template.add_run(' mit ').italic = True
    template.add_run('[GENAUIGKEIT/UNTERSTÜTZUNG]').bold = True
    template.add_run(' wie gemessen durch ').italic = True
    template.add_run('[BEWERTUNGSMETHODE]').bold = True
    
    # Sample goals
    doc.add_heading('Beispielziele', 2)
    
    doc.add_heading('Leseverstehen', 3)
    goals = [
        "Bis Mai 2026 wird [Schüler/in] Texte auf Klassenstufenniveau lesen und Verständnisfragen mit 80% Genauigkeit in 4 von 5 Versuchen beantworten.",
        "Bis Dezember 2025 wird [Schüler/in] Hauptideen und drei unterstützende Details in Sachtexten mit 85% Genauigkeit identifizieren.",
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_heading('Mathematik', 3)
    goals = [
        "Bis Juni 2026 wird [Schüler/in] mehrstufige Textaufgaben mit Addition und Subtraktion mit 80% Genauigkeit in 4 von 5 Gelegenheiten lösen.",
        "Bis April 2026 wird [Schüler/in] Verständnis der Einmaleins-Reihen 0-12 durch Lösen von Aufgaben mit 90% Genauigkeit zeigen.",
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Erstellt von Zaza Draft - KI-Werkzeuge für Lehrkräfte')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    
    doc.save(f"{OUTPUT_DIR}/foerderplan-vorlagen.docx")
    print("   ✅ Complete\n")

# ============================================================================
# 5. ENGLISH LESSON PLAN TEMPLATES (Word)
# ============================================================================
def create_lesson_templates_en():
    print("Creating 5/14: lesson-plan-templates.docx (English)...")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('AI-Enhanced Lesson Plan Templates', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(139, 92, 246)
    
    # Introduction
    doc.add_heading('Standard Lesson Plan Template', 1)
    
    sections = [
        ('Lesson Title:', '___________________________________________________________________'),
        ('Grade Level:', '____________  Subject: ____________  Duration: ____________'),
        ('', ''),
    ]
    
    for label, line in sections:
        if label:
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            if line:
                p.add_run(' ' + line)
    
    # Learning Objectives
    doc.add_heading('Learning Objectives', 2)
    doc.add_paragraph('By the end of this lesson, students will be able to:')
    for i in range(3):
        doc.add_paragraph('___________________________________________________________________', style='List Bullet')
    
    # Standards
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Standards Addressed:').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    # Materials
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Materials Needed:').bold = True
    doc.add_paragraph('✓ AI Suggestion: Ask AI to "List materials needed for a [grade] lesson on [topic]"', style='List Bullet')
    for i in range(3):
        doc.add_paragraph('___________________________________________________________________', style='List Bullet')
    
    # Lesson Structure
    doc.add_heading('Lesson Structure', 2)
    
    p = doc.add_paragraph()
    p.add_run('1. Introduction/Hook (5-10 minutes)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    doc.add_paragraph('✓ AI Suggestion: "Create an engaging hook for a [grade] lesson about [topic]"', style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('2. Direct Instruction (15-20 minutes)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('3. Guided Practice (10-15 minutes)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('4. Independent Practice (15-20 minutes)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('5. Closure/Assessment (5 minutes)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Created by Zaza Draft - AI Tools for Teachers')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    
    doc.save(f"{OUTPUT_DIR}/lesson-plan-templates.docx")
    print("   ✅ Complete\n")

# ============================================================================
# 6. GERMAN LESSON PLAN TEMPLATES (Word)
# ============================================================================
def create_unterrichtsplan_templates_de():
    print("Creating 6/14: unterrichtsplan-vorlagen.docx (German)...")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('KI-unterstützte Unterrichtsplan-Vorlagen', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(139, 92, 246)
    
    # Introduction
    doc.add_heading('Standard-Unterrichtsplan-Vorlage', 1)
    
    sections = [
        ('Unterrichtsthema:', '___________________________________________________________________'),
        ('Klassenstufe:', '____________  Fach: ____________  Dauer: ____________'),
        ('', ''),
    ]
    
    for label, line in sections:
        if label:
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            if line:
                p.add_run(' ' + line)
    
    # Learning Objectives
    doc.add_heading('Lernziele', 2)
    doc.add_paragraph('Am Ende dieser Unterrichtsstunde werden die Schüler in der Lage sein:')
    for i in range(3):
        doc.add_paragraph('___________________________________________________________________', style='List Bullet')
    
    # Standards
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Bildungsstandards:').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    # Materials
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Benötigte Materialien:').bold = True
    doc.add_paragraph('✓ KI-Vorschlag: Fragen Sie die KI: "Liste Materialien für eine [Klassenstufe]-Stunde zum Thema [Thema]"', style='List Bullet')
    for i in range(3):
        doc.add_paragraph('___________________________________________________________________', style='List Bullet')
    
    # Lesson Structure
    doc.add_heading('Unterrichtsablauf', 2)
    
    p = doc.add_paragraph()
    p.add_run('1. Einstieg/Motivation (5-10 Minuten)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    doc.add_paragraph('✓ KI-Vorschlag: "Erstelle einen motivierenden Einstieg für eine [Klassenstufe]-Stunde zum Thema [Thema]"', style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('2. Erarbeitung (15-20 Minuten)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('3. Übungsphase (10-15 Minuten)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('4. Vertiefung/Anwendung (15-20 Minuten)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('5. Abschluss/Sicherung (5 Minuten)').bold = True
    doc.add_paragraph('___________________________________________________________________')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Erstellt von Zaza Draft - KI-Werkzeuge für Lehrkräfte')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    
    doc.save(f"{OUTPUT_DIR}/unterrichtsplan-vorlagen.docx")
    print("   ✅ Complete\n")

# ============================================================================
# HTML FILES (7-14): These need to be converted to PDF after creation
# ============================================================================

def create_html_file(filename, title, content):
    """Helper function to create HTML files with consistent styling"""
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        @page {{ margin: 2cm; }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        .header {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 3px solid #8B5CF6;
        }}
        .logo {{
            color: #8B5CF6;
            font-size: 28px;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        h1 {{
            color: #1a1a1a;
            font-size: 32px;
            margin: 20px 0 10px 0;
        }}
        h2 {{
            color: #8B5CF6;
            font-size: 24px;
            margin-top: 30px;
            margin-bottom: 15px;
            padding-bottom: 10px;
            border-bottom: 2px solid #e0e0e0;
        }}
        .footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 2px solid #e0e0e0;
            text-align: center;
            color: #666;
            font-size: 14px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="logo">🎓 Zaza Draft</div>
        <h1>{title}</h1>
    </div>
    
    {content}
    
    <div class="footer">
        <p><strong>Zaza Draft - AI Tools for Teachers</strong></p>
        <p>Visit zazadraft.com for more resources</p>
        <p>© 2025 Zaza Technologies. All rights reserved.</p>
    </div>
</body>
</html>"""
    
    with open(f"{OUTPUT_DIR}/{filename}", 'w', encoding='utf-8') as f:
        f.write(html)

def create_getting_started_en():
    print("Creating 7/14: getting-started-ai.html (English)...")
    
    content = """
    <h2>What is AI?</h2>
    <p><strong>Artificial Intelligence (AI)</strong> refers to computer systems that can perform tasks that typically require human intelligence—like understanding language, recognizing patterns, and solving problems.</p>
    
    <h2>Why Use AI in the Classroom?</h2>
    <ul>
        <li><strong>Save Time:</strong> Automate lesson planning, grading feedback, and administrative tasks</li>
        <li><strong>Personalize Learning:</strong> Generate differentiated materials for different student levels</li>
        <li><strong>Provide Instant Feedback:</strong> Students receive immediate, detailed feedback on their work</li>
        <li><strong>Enhance Creativity:</strong> Brainstorm lesson ideas and create engaging activities</li>
        <li><strong>Support Diverse Learners:</strong> Generate materials for ELL students, IEPs, and advanced learners</li>
    </ul>
    
    <h2>Top 5 AI Tools to Try First</h2>
    <p><strong>1. ChatGPT</strong> - Conversational AI for lesson planning and content creation</p>
    <p><strong>2. Zaza Draft</strong> - AI writing assistant specifically for teachers</p>
    <p><strong>3. Google Gemini</strong> - Integrated with Google Workspace</p>
    <p><strong>4. Canva AI</strong> - Create visual materials and presentations</p>
    <p><strong>5. Diffit</strong> - Adapt reading materials to different levels</p>
    
    <h2>5 Quick Wins This Week</h2>
    <p><strong>Quick Win #1:</strong> Generate a week of bell ringers</p>
    <p><strong>Quick Win #2:</strong> Differentiate a reading passage into 3 levels</p>
    <p><strong>Quick Win #3:</strong> Draft parent email responses</p>
    <p><strong>Quick Win #4:</strong> Create practice problems for math</p>
    <p><strong>Quick Win #5:</strong> Summarize an article for students</p>
    
    <h2>Common Concerns Addressed</h2>
    <p><strong>Will AI replace teachers?</strong> No. AI is a tool to enhance your capabilities, not replace them.</p>
    <p><strong>Student privacy?</strong> Always use FERPA-compliant tools and never input personally identifiable information.</p>
    <p><strong>Not tech-savvy enough?</strong> If you can type a question, you can use AI. No coding required.</p>
    """
    
    create_html_file("getting-started-ai.html", "Getting Started with AI", content)
    print("   ✅ Complete (convert to PDF)\n")

def create_getting_started_de():
    print("Creating 8/14: ki-einstieg.html (German)...")
    
    content = """
    <h2>Was ist KI?</h2>
    <p><strong>Künstliche Intelligenz (KI)</strong> bezeichnet Computersysteme, die Aufgaben ausführen können, die normalerweise menschliche Intelligenz erfordern—wie Sprache verstehen, Muster erkennen und Probleme lösen.</p>
    
    <h2>Warum KI im Unterricht einsetzen?</h2>
    <ul>
        <li><strong>Zeit sparen:</strong> Automatisieren Sie Unterrichtsplanung, Feedback und administrative Aufgaben</li>
        <li><strong>Lernen personalisieren:</strong> Generieren Sie differenzierte Materialien für verschiedene Schülerniveaus</li>
        <li><strong>Sofortiges Feedback:</strong> Schüler erhalten unmittelbare, detaillierte Rückmeldungen</li>
        <li><strong>Kreativität fördern:</strong> Brainstormen Sie Unterrichtsideen und erstellen Sie ansprechende Aktivitäten</li>
        <li><strong>Vielfalt unterstützen:</strong> Erstellen Sie Materialien für DaZ-Schüler, Förderpläne und leistungsstarke Lernende</li>
    </ul>
    
    <h2>Top 5 KI-Tools für den Einstieg</h2>
    <p><strong>1. ChatGPT</strong> - Konversations-KI für Unterrichtsplanung und Inhaltserstellung</p>
    <p><strong>2. Zaza Draft</strong> - KI-Schreibassistent speziell für Lehrkräfte</p>
    <p><strong>3. Google Gemini</strong> - Integriert mit Google Workspace</p>
    <p><strong>4. Canva AI</strong> - Erstellen Sie visuelle Materialien und Präsentationen</p>
    <p><strong>5. Diffit</strong> - Passen Sie Lesetexte an verschiedene Niveaus an</p>
    
    <h2>5 Schnelle Erfolge diese Woche</h2>
    <p><strong>Schneller Erfolg #1:</strong> Generieren Sie eine Woche Einstiegsaufgaben</p>
    <p><strong>Schneller Erfolg #2:</strong> Differenzieren Sie einen Lesetext in 3 Niveaus</p>
    <p><strong>Schneller Erfolg #3:</strong> Entwerfen Sie Eltern-E-Mail-Antworten</p>
    <p><strong>Schneller Erfolg #4:</strong> Erstellen Sie Übungsaufgaben für Mathematik</p>
    <p><strong>Schneller Erfolg #5:</strong> Fassen Sie einen Artikel für Schüler zusammen</p>
    
    <h2>Häufige Bedenken angesprochen</h2>
    <p><strong>Wird KI Lehrer ersetzen?</strong> Nein. KI ist ein Werkzeug zur Verbesserung Ihrer Fähigkeiten, nicht zum Ersetzen.</p>
    <p><strong>Datenschutz für Schüler?</strong> Verwenden Sie immer DSGVO-konforme Tools und geben Sie niemals personenbezogene Daten ein.</p>
    <p><strong>Nicht technisch versiert genug?</strong> Wenn Sie eine Frage tippen können, können Sie KI verwenden. Keine Programmierung erforderlich.</p>
    """
    
    create_html_file("ki-einstieg.html", "Einstieg in die KI", content)
    print("   ✅ Complete (convert to PDF)\n")

def create_tool_evaluation_en():
    print("Creating 9/14: ai-tool-evaluation.html (English)...")
    
    content = """
    <h2>Evaluation Categories</h2>
    
    <h3>1. Privacy & Security</h3>
    <p>☐ FERPA Compliance - Tool complies with student privacy laws</p>
    <p>☐ Data Storage - Clear policies on data storage and handling</p>
    <p>☐ Data Sharing - Does not sell or share student data</p>
    <p>☐ User Authentication - Secure login and access controls</p>
    
    <h3>2. Pedagogical Value</h3>
    <p>☐ Learning Goals - Supports meaningful educational objectives</p>
    <p>☐ Age Appropriate - Content suitable for target age group</p>
    <p>☐ Differentiation - Accommodates diverse learning needs</p>
    <p>☐ Research-Based - Grounded in educational research</p>
    
    <h3>3. Ease of Use</h3>
    <p>☐ Teacher Learning Curve - Easy for teachers to learn</p>
    <p>☐ Student Accessibility - Students can use independently</p>
    <p>☐ Interface Design - Clean and intuitive interface</p>
    <p>☐ Multi-Platform - Works across devices</p>
    
    <h3>4. Cost & Sustainability</h3>
    <p>☐ Free Version - Free tier provides meaningful functionality</p>
    <p>☐ Pricing Transparency - Clear pricing structure</p>
    <p>☐ Value for Money - Benefits justify the cost</p>
    <p>☐ Budget Feasibility - Fits typical school budgets</p>
    
    <h3>5. Support & Resources</h3>
    <p>☐ Training - Tutorials and guides available</p>
    <p>☐ Customer Support - Responsive customer service</p>
    <p>☐ User Community - Active community for support</p>
    <p>☐ Documentation - Clear, comprehensive documentation</p>
    
    <h2>Scoring Guide</h2>
    <p>Rate each item 1-5:</p>
    <p>5 = Excellent | 4 = Good | 3 = Acceptable | 2 = Poor | 1 = Unacceptable</p>
    """
    
    create_html_file("ai-tool-evaluation.html", "AI Tool Evaluation Checklist", content)
    print("   ✅ Complete (convert to PDF)\n")

def create_tool_evaluation_de():
    print("Creating 10/14: ki-tool-bewertung.html (German)...")
    
    content = """
    <h2>Bewertungskategorien</h2>
    
    <h3>1. Datenschutz & Sicherheit</h3>
    <p>☐ DSGVO-Konformität - Tool entspricht Datenschutzgesetzen</p>
    <p>☐ Datenspeicherung - Klare Richtlinien zur Datenspeicherung</p>
    <p>☐ Datenweitergabe - Verkauft oder teilt keine Schülerdaten</p>
    <p>☐ Benutzerauthentifizierung - Sichere Anmeldung und Zugriffskontrollen</p>
    
    <h3>2. Pädagogischer Wert</h3>
    <p>☐ Lernziele - Unterstützt sinnvolle Bildungsziele</p>
    <p>☐ Altersgerecht - Inhalte passend für Zielgruppe</p>
    <p>☐ Differenzierung - Berücksichtigt diverse Lernbedürfnisse</p>
    <p>☐ Forschungsbasiert - Auf Bildungsforschung gegründet</p>
    
    <h3>3. Benutzerfreundlichkeit</h3>
    <p>☐ Lehrkraft-Lernkurve - Einfach für Lehrkräfte zu erlernen</p>
    <p>☐ Schülerzugang - Schüler können eigenständig nutzen</p>
    <p>☐ Interface-Design - Sauberes und intuitives Interface</p>
    <p>☐ Multiplattform - Funktioniert auf verschiedenen Geräten</p>
    
    <h3>4. Kosten & Nachhaltigkeit</h3>
    <p>☐ Kostenlose Version - Kostenlose Version bietet sinnvolle Funktionen</p>
    <p>☐ Preistransparenz - Klare Preisstruktur</p>
    <p>☐ Preis-Leistung - Nutzen rechtfertigt Kosten</p>
    <p>☐ Budget-Machbarkeit - Passt in typische Schulbudgets</p>
    
    <h3>5. Support & Ressourcen</h3>
    <p>☐ Schulung - Tutorials und Anleitungen verfügbar</p>
    <p>☐ Kundensupport - Reaktionsschneller Kundenservice</p>
    <p>☐ Benutzer-Community - Aktive Community für Unterstützung</p>
    <p>☐ Dokumentation - Klare, umfassende Dokumentation</p>
    
    <h2>Bewertungsskala</h2>
    <p>Bewerten Sie jedes Element 1-5:</p>
    <p>5 = Ausgezeichnet | 4 = Gut | 3 = Akzeptabel | 2 = Schwach | 1 = Inakzeptabel</p>
    """
    
    create_html_file("ki-tool-bewertung.html", "KI-Tool-Bewertungs-Checkliste", content)
    print("   ✅ Complete (convert to PDF)\n")

def create_ethics_framework_en():
    print("Creating 11/14: ai-ethics-framework.html (English)...")
    
    content = """
    <h2>The Five Core Ethical Principles</h2>
    
    <h3>1. Transparency</h3>
    <p><strong>Principle:</strong> Students and parents should know when and how AI is being used.</p>
    <p><strong>In Practice:</strong></p>
    <ul>
        <li>Clearly communicate when students are interacting with AI-generated content</li>
        <li>Inform parents about AI tools used in the classroom</li>
        <li>Explain how AI influences decisions</li>
    </ul>
    
    <h3>2. Fairness & Non-Discrimination</h3>
    <p><strong>Principle:</strong> AI tools should promote equity and avoid perpetuating bias.</p>
    <p><strong>In Practice:</strong></p>
    <ul>
        <li>Be aware that AI can reflect societal biases</li>
        <li>Review AI-generated content for stereotypes</li>
        <li>Ensure equitable access to AI learning opportunities</li>
    </ul>
    
    <h3>3. Privacy & Data Protection</h3>
    <p><strong>Principle:</strong> Student data must be protected and AI tools should minimize data collection.</p>
    <p><strong>In Practice:</strong></p>
    <ul>
        <li>Use only FERPA-compliant AI tools</li>
        <li>Never input personally identifiable information</li>
        <li>Understand what data AI tools collect</li>
    </ul>
    
    <h3>4. Accountability & Teacher Oversight</h3>
    <p><strong>Principle:</strong> Educators remain responsible for educational decisions.</p>
    <p><strong>In Practice:</strong></p>
    <ul>
        <li>Always review AI-generated content before using with students</li>
        <li>Make final grading decisions yourself</li>
        <li>Verify factual accuracy of AI information</li>
    </ul>
    
    <h3>5. Human-Centeredness</h3>
    <p><strong>Principle:</strong> AI should enhance human capabilities, not replace essential human elements.</p>
    <p><strong>In Practice:</strong></p>
    <ul>
        <li>Use AI to save time for more meaningful interactions</li>
        <li>Prioritize human judgment in complex decisions</li>
        <li>Maintain teacher-student relationships as central</li>
    </ul>
    
    <h2>Ethical Decision-Making Flowchart</h2>
    <p>Before using AI, ask yourself:</p>
    <ul>
        <li>✓ Does this tool comply with student privacy laws?</li>
        <li>✓ Will students/parents know AI is being used?</li>
        <li>✓ Have I reviewed the output for bias and accuracy?</li>
        <li>✓ Am I using AI to enhance (not replace) my judgment?</li>
        <li>✓ Does this prioritize student learning and wellbeing?</li>
    </ul>
    """
    
    create_html_file("ai-ethics-framework.html", "AI Ethics Framework for Educators", content)
    print("   ✅ Complete (convert to PDF)\n")

def create_ethics_framework_de():
    print("Creating 12/14: ki-ethik-rahmen.html (German)...")
    
    content = """
    <h2>Die fünf ethischen Kernprinzipien</h2>
    
    <h3>1. Transparenz</h3>
    <p><strong>Prinzip:</strong> Schüler und Eltern sollten wissen, wann und wie KI eingesetzt wird.</p>
    <p><strong>In der Praxis:</strong></p>
    <ul>
        <li>Kommunizieren Sie klar, wenn Schüler mit KI-generierten Inhalten interagieren</li>
        <li>Informieren Sie Eltern über im Unterricht verwendete KI-Tools</li>
        <li>Erklären Sie, wie KI Entscheidungen beeinflusst</li>
    </ul>
    
    <h3>2. Fairness & Nicht-Diskriminierung</h3>
    <p><strong>Prinzip:</strong> KI-Tools sollten Gleichberechtigung fördern und Vorurteile vermeiden.</p>
    <p><strong>In der Praxis:</strong></p>
    <ul>
        <li>Seien Sie sich bewusst, dass KI gesellschaftliche Vorurteile widerspiegeln kann</li>
        <li>Überprüfen Sie KI-generierte Inhalte auf Stereotypen</li>
        <li>Gewährleisten Sie gleichberechtigten Zugang zu KI-Lernmöglichkeiten</li>
    </ul>
    
    <h3>3. Datenschutz & Datensicherheit</h3>
    <p><strong>Prinzip:</strong> Schülerdaten müssen geschützt werden und KI-Tools sollten Datenerfassung minimieren.</p>
    <p><strong>In der Praxis:</strong></p>
    <ul>
        <li>Verwenden Sie nur DSGVO-konforme KI-Tools</li>
        <li>Geben Sie niemals personenbezogene Daten ein</li>
        <li>Verstehen Sie, welche Daten KI-Tools sammeln</li>
    </ul>
    
    <h3>4. Verantwortlichkeit & Lehreraufsicht</h3>
    <p><strong>Prinzip:</strong> Lehrkräfte bleiben für Bildungsentscheidungen verantwortlich.</p>
    <p><strong>In der Praxis:</strong></p>
    <ul>
        <li>Überprüfen Sie immer KI-generierte Inhalte vor der Verwendung mit Schülern</li>
        <li>Treffen Sie Bewertungsentscheidungen selbst</li>
        <li>Überprüfen Sie die sachliche Richtigkeit von KI-Informationen</li>
    </ul>
    
    <h3>5. Menschenzentriertheit</h3>
    <p><strong>Prinzip:</strong> KI sollte menschliche Fähigkeiten verbessern, nicht wesentliche menschliche Elemente ersetzen.</p>
    <p><strong>In der Praxis:</strong></p>
    <ul>
        <li>Nutzen Sie KI, um Zeit für bedeutungsvollere Interaktionen zu sparen</li>
        <li>Priorisieren Sie menschliches Urteilsvermögen bei komplexen Entscheidungen</li>
        <li>Behalten Sie Lehrer-Schüler-Beziehungen als zentral bei</li>
    </ul>
    
    <h2>Ethisches Entscheidungs-Flussdiagramm</h2>
    <p>Bevor Sie KI verwenden, fragen Sie sich:</p>
    <ul>
        <li>✓ Entspricht dieses Tool den Datenschutzgesetzen für Schüler?</li>
        <li>✓ Werden Schüler/Eltern wissen, dass KI verwendet wird?</li>
        <li>✓ Habe ich die Ausgabe auf Voreingenommenheit und Richtigkeit überprüft?</li>
        <li>✓ Verwende ich KI zur Verbesserung (nicht zum Ersetzen) meines Urteils?</li>
        <li>✓ Priorisiert dies das Lernen und Wohlbefinden der Schüler?</li>
    </ul>
    """
    
    create_html_file("ki-ethik-rahmen.html", "KI-Ethik-Rahmen für Lehrkräfte", content)
    print("   ✅ Complete (convert to PDF)\n")

def create_parent_emails_en():
    print("Creating 13/14: parent-email-templates.html (English)...")
    
    content = """
    <h2>Template 1: Introduction to AI in Classroom</h2>
    <p><strong>Subject:</strong> Exciting Updates About AI-Enhanced Learning</p>
    <p>Dear [Parent/Guardian Name],</p>
    <p>I'm writing to share some exciting updates about how we're enhancing learning in [your classroom] this year. We'll be thoughtfully integrating AI tools to support your child's learning.</p>
    <p><strong>These tools will help us:</strong></p>
    <ul>
        <li>Provide more personalized learning experiences</li>
        <li>Give faster, more detailed feedback on assignments</li>
        <li>Create engaging, differentiated activities</li>
        <li>Save time on administrative tasks</li>
    </ul>
    <p>Your child will still receive the same personal attention and direct instruction. AI is simply a tool that helps me be more effective.</p>
    <p>Please feel free to reach out with any questions.</p>
    <p>Best regards,<br>[Your Name]</p>
    
    <h2>Template 2: Monthly Progress Update</h2>
    <p><strong>Subject:</strong> [Student Name]'s Progress Update</p>
    <p>Dear [Parent/Guardian Name],</p>
    <p>I wanted to share [Student Name]'s progress this month and highlight some wonderful achievements.</p>
    <p><strong>Academic Highlights:</strong></p>
    <ul>
        <li>[Specific achievement or improvement area]</li>
        <li>[Growth observed in particular subject/skill]</li>
    </ul>
    <p><strong>How AI Tools Are Supporting Learning:</strong></p>
    <p>This month, [Student Name] has been using AI-assisted tools for [specific activity]. These tools have helped by [specific benefit].</p>
    <p>Please let me know if you'd like to discuss further.</p>
    <p>Best regards,<br>[Your Name]</p>
    
    <h2>Template 3: Addressing Concerns</h2>
    <p><strong>Subject:</strong> Re: Your Questions About AI</p>
    <p>Dear [Parent/Guardian Name],</p>
    <p>Thank you for reaching out with your questions about AI use in our classroom. I appreciate your engagement.</p>
    <p><strong>Your Concern:</strong> [Summarize parent's concern]</p>
    <p><strong>My Response:</strong> [Provide detailed, empathetic response]</p>
    <p>I value your perspective and want to ensure [Student Name] has the best possible learning experience. Would you be available for a call to discuss this further?</p>
    <p>Sincerely,<br>[Your Name]</p>
    """
    
    create_html_file("parent-email-templates.html", "Parent Email Templates", content)
    print("   ✅ Complete (convert to PDF)\n")

def create_parent_emails_de():
    print("Creating 14/14: eltern-email-vorlagen.html (German)...")
    
    content = """
    <h2>Vorlage 1: Einführung der KI im Unterricht</h2>
    <p><strong>Betreff:</strong> Spannende Neuigkeiten zum KI-unterstützten Lernen</p>
    <p>Sehr geehrte/r [Name des Erziehungsberechtigten],</p>
    <p>ich möchte Ihnen einige spannende Neuigkeiten mitteilen, wie wir das Lernen in [Ihrer Klasse] in diesem Jahr verbessern. Wir werden KI-Tools durchdacht integrieren, um das Lernen Ihres Kindes zu unterstützen.</p>
    <p><strong>Diese Tools helfen uns:</strong></p>
    <ul>
        <li>Personalisiertere Lernerfahrungen bereitzustellen</li>
        <li>Schnelleres, detaillierteres Feedback zu Aufgaben zu geben</li>
        <li>Ansprechende, differenzierte Aktivitäten zu erstellen</li>
        <li>Zeit bei administrativen Aufgaben zu sparen</li>
    </ul>
    <p>Ihr Kind wird weiterhin die gleiche persönliche Aufmerksamkeit und direkte Anleitung erhalten. KI ist einfach ein Werkzeug, das mir hilft, effektiver zu sein.</p>
    <p>Bei Fragen können Sie sich gerne an mich wenden.</p>
    <p>Mit freundlichen Grüßen,<br>[Ihr Name]</p>
    
    <h2>Vorlage 2: Monatlicher Fortschrittsbericht</h2>
    <p><strong>Betreff:</strong> Fortschrittsbericht für [Schülername]</p>
    <p>Sehr geehrte/r [Name des Erziehungsberechtigten],</p>
    <p>Ich möchte Ihnen den Fortschritt von [Schülername] in diesem Monat mitteilen und einige wunderbare Erfolge hervorheben.</p>
    <p><strong>Akademische Höhepunkte:</strong></p>
    <ul>
        <li>[Spezifische Leistung oder Verbesserungsbereich]</li>
        <li>[Beobachtetes Wachstum in bestimmtem Fach/Fähigkeit]</li>
    </ul>
    <p><strong>Wie KI-Tools das Lernen unterstützen:</strong></p>
    <p>In diesem Monat hat [Schülername] KI-unterstützte Tools für [spezifische Aktivität] verwendet. Diese Tools haben geholfen, indem [spezifischer Nutzen].</p>
    <p>Lassen Sie mich wissen, wenn Sie weitere Informationen wünschen.</p>
    <p>Mit freundlichen Grüßen,<br>[Ihr Name]</p>
    
    <h2>Vorlage 3: Bedenken ansprechen</h2>
    <p><strong>Betreff:</strong> Ihre Fragen zur KI</p>
    <p>Sehr geehrte/r [Name des Erziehungsberechtigten],</p>
    <p>Vielen Dank, dass Sie sich mit Ihren Fragen zur KI-Nutzung in unserem Unterricht an mich gewandt haben. Ich schätze Ihr Engagement.</p>
    <p><strong>Ihre Bedenken:</strong> [Zusammenfassung der Elternbedenken]</p>
    <p><strong>Meine Antwort:</strong> [Detaillierte, einfühlsame Antwort]</p>
    <p>Ich schätze Ihre Perspektive und möchte sicherstellen, dass [Schülername] die bestmögliche Lernerfahrung hat. Wären Sie für ein Gespräch verfügbar?</p>
    <p>Mit freundlichen Grüßen,<br>[Ihr Name]</p>
    """
    
    create_html_file("eltern-email-vorlagen.html", "Eltern-E-Mail-Vorlagen", content)
    print("   ✅ Complete (convert to PDF)\n")

# ============================================================================
# MAIN EXECUTION
# ============================================================================

if __name__ == "__main__":
    try:
        # Create all Excel files
        create_report_card_comments_en()
        create_report_card_comments_de()
        
        # Create all Word documents
        create_iep_templates_en()
        create_foerderplan_templates_de()
        create_lesson_templates_en()
        create_unterrichtsplan_templates_de()
        
        # Create all HTML files
        create_getting_started_en()
        create_getting_started_de()
        create_tool_evaluation_en()
        create_tool_evaluation_de()
        create_ethics_framework_en()
        create_ethics_framework_de()
        create_parent_emails_en()
        create_parent_emails_de()
        
        print("=" * 70)
        print("✅ SUCCESS! All 14 resources created")
        print("=" * 70)
        print(f"\nFiles location: {OUTPUT_DIR}/")
        print("\nNext steps:")
        print("1. Convert the 8 HTML files to PDF (Ctrl+P → Save as PDF in browser)")
        print("2. Move all files to: C:\\Users\\User\\zaza-draft-website-v0-3g\\public\\downloads\\")
        print("3. Deploy with: git add . && git commit -m 'feat: add all resources' && git push")
        print("\n✨ Your Zaza Draft resources are ready!")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("Make sure you have installed: pip install openpyxl python-docx")